#!/usr/bin/env python3
"""
Convert a ZyBooks-exported zip into a YAML multiple-choice question bank.

Supported zip layouts
---------------------
1. **QTI XML** (preferred): an ``items/`` folder containing one QTI v2.1
   XML file per question.
2. **Word documents** (legacy fallback): three ``.docx`` files following the
   ZyBooks naming convention.

Usage:
    python zybooks_zip_to_yaml.py export.zip output.yaml

Requirements:
    pip install python-docx pyyaml    (Word path only)
"""

import sys
import re
import shutil
import zipfile
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path

import yaml


# ---------- Utility: detect files in extracted zip ----------

def find_docs_in_extracted_dir(extract_dir: Path):
    """Return (questions_doc_path, answer_key_doc_path) from an extracted ZyBooks zip."""
    docx_files = list(extract_dir.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError("No .docx files found in the extracted zip directory.")

    questions_doc = None
    answer_key_doc = None

    # Heuristics based on ZyBooks naming conventions
    for p in docx_files:
        name_lower = p.name.lower()
        if "answer_key" in name_lower:
            answer_key_doc = p
        elif "with_answers" in name_lower:
            # We typically don't need this; ignore
            continue
        else:
            # Likely the base questions file (e.g., Test_Test.docx)
            questions_doc = p

    if questions_doc is None:
        raise FileNotFoundError("Could not find the questions .docx (without 'answer_key' or 'with_answers' in the name).")

    if answer_key_doc is None:
        raise FileNotFoundError("Could not find the answer key .docx (name containing 'answer_key').")

    return questions_doc, answer_key_doc


# ---------- Parsing answer key ----------

def parse_answer_key(answer_doc_path: Path):
    """
    Parse an answer-key .docx where each answer is on a line like:
        '1) b'
        '2) a'
    Returns: dict[int -> 'a'..'d']
    """
    from docx import Document
    doc = Document(str(answer_doc_path))
    answers = {}

    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r'^(\d+)\)\s*([a-dA-D])', text)
        if m:
            qnum = int(m.group(1))
            letter = m.group(2).lower()
            answers[qnum] = letter

    return answers


# ---------- Parsing questions ----------

def paragraph_text(paragraph):
    """Extract text from a paragraph, inserting spaces between runs at font boundaries.

    Word splits text into "runs" at every formatting change.  ``paragraph.text``
    simply concatenates the run texts, which drops whitespace when a font change
    falls between two words (e.g. "Hello" in bold + "world" in normal becomes
    "Helloworld").  This helper inserts a single space between consecutive runs
    when neither side already has whitespace at the boundary.
    """
    runs = paragraph.runs
    if not runs:
        return paragraph.text  # fall back for edge cases (e.g. field codes)
    parts = []
    for i, run in enumerate(runs):
        t = run.text
        if i > 0 and t and parts:
            prev = parts[-1]
            if prev and not prev[-1].isspace() and not t[0].isspace():
                parts.append(" ")
        parts.append(t)
    return "".join(parts)


def is_question_start(text: str) -> bool:
    return re.match(r'^\d+\)\s', text.strip()) is not None


def is_choice_line(text: str) -> bool:
    return re.match(r'^[a-dA-D]\.\s', text.strip()) is not None


def parse_questions(questions_doc_path: Path, answers: dict):
    """
    Parse the questions .docx into a list of dicts matching the YAML schema:

    - id: Q1
      points: 1
      type: mcq
      stem: [ {type: text|code, text: "...", ...}, ... ]
      choices: [ {key: 'a', text: '...'}, ... ]
      correct: 'b'
    """
    from docx import Document
    doc = Document(str(questions_doc_path))
    paras = [paragraph_text(p) for p in doc.paragraphs]
    n = len(paras)
    i = 0
    questions = []

    # Skip title or other non-question paragraphs at top
    while i < n and not is_question_start(paras[i]):
        i += 1

    while i < n:
        line = paras[i].strip()
        if not is_question_start(line):
            i += 1
            continue

        # Parse "N) ..." line
        m = re.match(r'^(\d+)\)\s*(.*)$', line)
        if not m:
            i += 1
            continue

        qnum = int(m.group(1))
        first_rest = m.group(2).strip()

        stem_blocks = []
        choices = []

        # If the first line has text after "N)", treat as stem text
        if first_rest:
            stem_blocks.append({
                "type": "text",
                "text": first_rest,
            })

        i += 1

        # Collect stem paragraphs until we hit choices or next question
        while i < n:
            raw = paras[i]
            text = raw.strip()

            # Skip blank lines
            if not text:
                i += 1
                continue

            # Next question?
            if is_question_start(text):
                break

            # Choices start?
            if is_choice_line(text):
                break

            print(raw)
            # Otherwise, part of stem
            if "\n" in raw:
                # Treat as code block with literal newlines
                stem_blocks.append({
                    "type": "code",
                    "language": "python",   # assumption for ENGR 131
                    "style": "mypython",    # matches your LaTeX listings style
                    "text": raw,
                })
            else:
                stem_blocks.append({
                    "type": "text",
                    "text": text,
                })

            i += 1

        # Collect choices if present (choices may span multiple paragraphs)
        while i < n:
            text = paras[i].strip()
            if not text:
                i += 1
                continue
            if is_question_start(text):
                break
            m = re.match(r'^([a-dA-D])\.\s*(.*)$', text)
            if m:
                # Start a new choice
                choices.append({
                    "key": m.group(1).lower(),
                    "text": m.group(2),
                })
            elif choices:
                # Continuation line for the current choice
                choices[-1]["text"] += " " + text
            else:
                # Non-choice text before any choice was seen — stop
                break
            i += 1

        # Build question dict
        qid = f"Q{qnum}"
        correct = answers.get(qnum)

        question = {
            "id": qid,
            "points": 1,
            "type": "mcq",
            "stem": stem_blocks,
            "choices": choices,
            "correct": correct,
        }

        questions.append(question)

        # After choices, loop continues; top of while will detect next question start

    return questions


# ---------- QTI v2.1 XML parsing ----------

_QTI_NS = "http://www.imsglobal.org/xsd/imsqti_v2p1"


def _qti(tag):
    """Return a namespace-qualified tag name for QTI v2.1."""
    return f"{{{_QTI_NS}}}{tag}"


def _collapse_blank_lines(text):
    """Collapse runs of multiple blank lines into a single blank line."""
    return re.sub(r'\n{3,}', '\n\n', text)


_TOPIC_INDICATORS = {
    "strings": [
        # stem/choice keywords
        r'\bstring\b', r'\bsubstring\b', r'\bcharacter\b', r'\bconcat',
        r'\bslice\b', r'\bslicing\b', r'\bf-string\b', r'\bformat',
        r'\bformatting\b', r'\bpresentation\s+type',
        # string methods
        r'\.split\b', r'\.join\b', r'\.strip\b', r'\.rstrip\b', r'\.lstrip\b',
        r'\.upper\b', r'\.lower\b', r'\.capitalize\b', r'\.title\b',
        r'\.find\b', r'\.rfind\b', r'\.index\b', r'\.count\b',
        r'\.replace\b', r'\.startswith\b', r'\.endswith\b',
        r'\.isdigit\b', r'\.isalpha\b', r'\.isalnum\b',
        r'\.isupper\b', r'\.islower\b', r'\.isspace\b',
        r'\.format\b', r'\.encode\b', r'\.decode\b',
        # string operations
        r'\[\s*-?\d*\s*:\s*-?\d*\s*(:\s*-?\d*\s*)?\]',  # slicing [0:6], [-5:], [0:-1:2]
        # format specifiers (alignment/fill)
        r'\{[^}]*:[^}]*[<>^][^}]*\}',   # {x:>4}, {x:*<4}
        # f-string / format() calls
        r'\bf"[^"]*"', r'\bf\'[^\']*\'',
    ],
    "functions": [
        r'\bfunction',  # matches "function" and "functions"
        r'\bparameter\b', r'\bargument\b',
        r'\bdef\s+\w+', r'\breturn\b',
        r'\bscope\b', r'\blocal\b', r'\bglobal\b',
        r'\blambda\b', r'\brecursion\b', r'\brecursive\b',
        r'\bdocstring',
    ],
    "files": [
        r'\bfile\b', r'\bfilename\b',
        r'\bopen\s*\(', r'\.read\b', r'\.readline\b', r'\.readlines\b',
        r'\.write\b', r'\.writelines\b', r'\.close\b',
        r'\bwith\s+open\b',
        r"['\"][rwa]b?['\"]",  # file modes 'r', 'w', 'a', 'rb'
        r'\bimport\s+csv\b', r'\bimport\s+json\b',
        r'\bcsv\.\w+', r'\bjson\.\w+',
        r'\bflush\b', r'\bbuffer\b',
        r'\bbytes\s*\(', r'\bbytes\s+literal',  # bytes() calls
    ],
    "lists": [
        r'\blist\b', r'\blist\s*\(',
        r'\.append\b', r'\.extend\b', r'\.insert\b',
        r'\.pop\b', r'\.remove\b', r'\.sort\b', r'\.reverse\b',
        r'\bsorted\b', r'\blen\b',
        r'\bfor\b.*\bin\b',  # iteration
        r'\[.*\bfor\b.*\bin\b',  # list comprehension
        r'\ball\s*\(', r'\bany\s*\(',  # list aggregate builtins
        r'\bmin\s*\(', r'\bmax\s*\(',
        r'\w_list\b',  # variable names like new_list, my_list
    ],
    "dictionaries": [
        r'\bdict\b', r'\bdictionar',
        r'\.keys\b', r'\.values\b', r'\.items\b',
        r'\.get\b', r'\.update\b',
        r'\{\s*["\'\w]+\s*:\s*[^<>=*]',  # dict literal {key: val} (exclude format specs)
    ],
    "modules": [
        r'\bmodule\b', r'\bpackage\b', r'\blibrary\b',
        r'\bimport\s+\w+', r'\bfrom\s+\w+\s+import\b',
        r'\bsys\b', r'\bos\b', r'\bmath\b',
        r'\bsys\.argv\b', r'\b__name__\b',
        r'\bcommand.line\s+argument',
        r'\bpip\b',
        r'\brandom\.\w+',  # random.randrange, random.randint
        r'\bnp\.\w+', r'\bnumpy\b', r'\bpandas\b',  # numpy/pandas
        r'\bstruct\.\w+',  # struct.pack, struct.unpack
    ],
    "classes": [
        r'\bclass\s+\w+', r'\bself\.',
        r'\b__init__\b', r'\binherit', r'\bsubclass\b',
        r'\bobject\b', r'\binstance\b', r'\bconstructor\b',
        r'\battribute\b', r'\bmethod\b',
        r'\bsuper\s*\(',
        r'\b__\w+__\s*\(',  # dunder method calls __mul__(, __mod__(
        r'\bisinstance\s*\(',
        r'\boverload',
    ],
}


def _detect_topic(stem_blocks, choices):
    """Classify a question into a topic based on stem and choice content."""
    # Combine all textual content for scanning
    parts = []
    for b in stem_blocks:
        parts.append(b.get("text", ""))
    for c in choices:
        parts.append(c.get("text", ""))
    blob = "\n".join(parts)

    scores = {}
    for topic, patterns in _TOPIC_INDICATORS.items():
        scores[topic] = sum(
            len(re.findall(p, blob, re.IGNORECASE)) for p in patterns
        )

    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "other"


def _looks_like_code(text):
    """Heuristic: return True if *text* looks like a code expression.

    Checks for common Python syntax indicators such as format specifiers,
    method calls, brackets, assignment operators, etc.
    """
    indicators = [
        r'\{.*[:!].*\}',       # format specifiers  {x:>4}, {x!r}
        r'\w+\.\w+\(',         # method calls        foo.bar(
        r'\w+\(',              # function calls       upper(), len(
        r'\w+\[',              # subscript            arr[
        r'\w+\s*[=!<>]=',      # comparison/assign    x == , x !=
        r'\breturn\b|\bprint\b|\bimport\b',                     # unambiguous keywords
        r'\bTrue\b|\bFalse\b|\bNone\b',                         # builtins
        # compound-statement keywords only when they look like Python
        # statements (line ends with a colon)
        r'(?m)^\s*(?:class|def|if|elif|else|for|while|with|try|except|finally)\b.*:\s*$',
    ]
    return any(re.search(p, text) for p in indicators)


def _wrap_inline_code(text):
    """Wrap string literals and Python identifiers in ``backtick`` markers.

    Detects double-quoted string literals and underscore-containing words
    (e.g. ``new_string``, ``my_list``) that are almost certainly Python
    identifiers.  Skips regions already inside ``...`` backtick markers.
    """
    # Split on ``...`` regions, only transform the outside parts
    parts = re.split(r'(``.*?``)', text)
    for i, part in enumerate(parts):
        if not part.startswith('``'):
            # Wrap double-quoted string literals
            part = re.sub(r'"([^"]+)"', r'``"\1"``', part)
            # Wrap underscore identifiers (e.g. new_string, my_list)
            part = re.sub(r'\b(\w+_\w+)\b', r'``\1``', part)
            parts[i] = part
    return "".join(parts)


def _table_to_text(table_elem):
    """Convert a QTI ``<table>`` element into aligned plain-text rows.

    Returns a string that resembles ``pandas`` DataFrame output, suitable
    for inclusion as a ``code`` stem block.
    """
    ns = _QTI_NS
    rows = []
    for tr in table_elem.iter(f"{{{ns}}}tr"):
        cells = []
        for td in tr.iter(f"{{{ns}}}td"):
            # A cell might contain nested <p> or bare text
            text = _element_text(td).strip()
            cells.append(text)
        rows.append(cells)
    if not rows:
        return ""

    # Pad all rows to the same width
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    # Column widths
    col_widths = [max(len(r[c]) for r in rows) for c in range(max_cols)]

    lines = []
    for r in rows:
        line = "  ".join(r[c].rjust(col_widths[c]) for c in range(max_cols))
        lines.append(line)
    return "\n".join(lines)


def _element_text(elem, *, inline_code=False):
    """Recursively extract text from an XML element, converting <br/> to newlines.

    Parameters
    ----------
    inline_code : bool
        If True, wrap ``<code>`` content with double-backtick markers
        (``...``) so the downstream LaTeX renderer can typeset them
        with ``\\inl{}``.
    """
    _BLOCK_TAGS = {"p", "div", "pre"}
    parts = []
    if elem.text:
        parts.append(elem.text)
    for child in elem:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "br":
            parts.append("\n")
        elif local == "code" and inline_code:
            code_text = _element_text(child)
            parts.append(f"``{code_text.strip()}``")
        else:
            # Insert a newline before block-level elements when content precedes them
            if local in _BLOCK_TAGS and parts and not parts[-1].endswith("\n"):
                parts.append("\n")
            parts.append(_element_text(child, inline_code=inline_code))
        if child.tail:
            parts.append(child.tail)
    return "".join(parts)


def parse_qti_item(xml_path):
    """Parse a single QTI v2.1 assessmentItem XML file into a question dict.

    Returns None if the file is not a valid assessmentItem (e.g. manifest).
    """
    tree = ET.parse(xml_path)
    root = tree.getroot()

    # Skip files that aren't assessmentItem elements
    root_local = root.tag.split("}")[-1] if "}" in root.tag else root.tag
    if root_local != "assessmentItem":
        return None

    qid = root.get("identifier", Path(xml_path).stem)

    # --- correct answer ---
    correct = None
    val = root.find(f"{_qti('responseDeclaration')}/{_qti('correctResponse')}/{_qti('value')}")
    if val is not None and val.text:
        correct = val.text.strip().lower()

    # --- item body ---
    item_body = root.find(_qti("itemBody"))
    if item_body is None:
        print(f"  WARNING: {Path(xml_path).name} has no itemBody, skipping")
        return None

    stem_blocks = []
    choices = []

    # Stem: <p> and <img> elements inside top-level <div> (not inside choiceInteraction)
    for div in item_body.findall(_qti("div")):
        for child in div:
            child_local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if child_local == "img":
                # Image element — record as image stem block
                src = child.get("src", "")
                alt = child.get("alt", "")
                width_attr = child.get("width", "")
                stem_blocks.append({
                    "type": "image",
                    "src": src,
                    "alt": alt,
                    "width": width_attr,
                })
                continue

            if child_local == "table":
                table_text = _table_to_text(child)
                if table_text:
                    stem_blocks.append({
                        "type": "code",
                        "language": "text",
                        "style": "mypython",
                        "text": table_text,
                    })
                continue

            if child_local != "p":
                continue

            p = child
            code = p.find(_qti("code"))
            if code is not None:
                # Distinguish block code (<p> is *only* a code listing)
                # from inline code (<code> mixed with surrounding text).
                has_surrounding_text = bool(
                    (p.text and p.text.strip())
                    or (code.tail and code.tail.strip())
                )
                if has_surrounding_text:
                    # Inline code — render as text with ``backtick`` markers
                    text = _wrap_inline_code(_element_text(p, inline_code=True).strip())
                    if text:
                        stem_blocks.append({"type": "text", "text": text})
                else:
                    # Block code listing
                    code_text = _collapse_blank_lines(_element_text(code).strip())
                    stem_blocks.append({
                        "type": "code",
                        "language": "python",
                        "style": "mypython",
                        "text": code_text,
                    })
            else:
                text = _wrap_inline_code(_element_text(p).strip())
                if text:
                    stem_blocks.append({"type": "text", "text": text})

    # Merge consecutive text blocks where the second is a sentence
    # continuation (ZyBooks splits <p> at inline-element boundaries).
    merged = []
    for block in stem_blocks:
        if (merged
                and merged[-1].get("type") == "text"
                and block.get("type") == "text"):
            text = block["text"]
            if text and (text[0].islower() or text[0] in ',.?!;:'):
                if text[0] in ',.?!;:':
                    merged[-1]["text"] += text
                else:
                    merged[-1]["text"] += " " + text
                continue
        merged.append(block)
    stem_blocks = merged

    # Determine if the stem implies choices are code/output
    stem_text_joined = " ".join(
        b["text"] for b in stem_blocks if b.get("type") == "text"
    ).lower()
    stem_implies_code = bool(re.search(
        r"what is (the )?(output|result|value)"
        r"|what does .* (print|output|return)"
        r"|what will .* (print|output|display)"
        r"|complete the (following )?code"
        r"|which .* (line|statement|expression|code|method)"
        r"|presentation type"
        r"|as output",
        stem_text_joined,
    ))

    # Choices
    ci = item_body.find(_qti("choiceInteraction"))
    if ci is not None:
        for sc in ci.findall(_qti("simpleChoice")):
            key = sc.get("identifier", "").lower()
            # Check if this individual choice contains <code>
            has_code_elem = sc.find(f".//{_qti('code')}") is not None
            text = _element_text(sc).strip()
            if has_code_elem:
                text = _collapse_blank_lines(text)
            choice = {"key": key, "text": text}
            if has_code_elem or stem_implies_code or _looks_like_code(text):
                choice["type"] = "code"
            choices.append(choice)

    topic = _detect_topic(stem_blocks, choices)

    return {
        "id": qid,
        "points": 1,
        "type": "mcq",
        "topic": topic,
        "stem": stem_blocks,
        "choices": choices,
        "correct": correct,
    }


def _numeric_sort_key(path):
    """Extract the trailing integer from a filename for numeric sorting.

    E.g. ``qti_item_VE_IP_42.xml`` -> 42.  Files without a number sort last.
    """
    m = re.search(r'(\d+)\.xml$', path.name)
    return int(m.group(1)) if m else float('inf')


def parse_qti_items_dir(items_dir):
    """Parse all QTI XML files in a directory, returning a numerically sorted list."""
    xml_files = sorted(items_dir.glob("*.xml"), key=_numeric_sort_key)
    if not xml_files:
        raise FileNotFoundError(f"No .xml files found in {items_dir}")
    questions = []
    for xf in xml_files:
        q = parse_qti_item(xf)
        if q is not None:
            questions.append(q)
    # Assign zero-padded sequential IDs
    width = len(str(len(questions)))
    for i, q in enumerate(questions, 1):
        q["id"] = f"Q{i:0{width}d}"
    print(f"Parsed {len(questions)} QTI items from {items_dir} ({len(xml_files) - len(questions)} skipped)")
    return questions


# ---------- Main entry points ----------

def convert(zip_path, output_yaml_path):
    """Convert a ZyBooks-exported zip into a YAML question bank.

    Auto-detects the format: if the zip contains an ``items/`` directory with
    XML files, the QTI parser is used.  Otherwise, falls back to the Word
    document parser.

    Parameters
    ----------
    zip_path : str or Path
        Path to the exported zip file.
    output_yaml_path : str or Path
        Path to the output YAML file.
    """
    zip_path = Path(zip_path)
    output_yaml_path = Path(output_yaml_path)

    if not zip_path.is_file():
        raise FileNotFoundError(f"{zip_path} does not exist or is not a file.")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)

        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(tmpdir_path)

        items_dir = tmpdir_path / "items"
        if items_dir.is_dir() and list(items_dir.glob("*.xml")):
            questions = parse_qti_items_dir(items_dir)
        else:
            questions_doc, answer_key_doc = find_docs_in_extracted_dir(tmpdir_path)
            print(f"Using questions doc:   {questions_doc.name}")
            print(f"Using answer key doc:  {answer_key_doc.name}")
            answers = parse_answer_key(answer_key_doc)
            questions = parse_questions(questions_doc, answers)

        # Collect image references and copy files alongside the YAML
        images_dir = output_yaml_path.parent / "images"
        image_count = 0
        for q in questions:
            for block in q.get("stem", []):
                if block.get("type") == "image":
                    src = block["src"]
                    src_path = items_dir / src if items_dir.is_dir() else None
                    if src_path and src_path.is_file():
                        images_dir.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(src_path, images_dir / src)
                        image_count += 1
                    else:
                        print(f"  WARNING: image {src} referenced but not found in zip")

        data = {"questions": questions}

        with output_yaml_path.open("w", encoding="utf-8") as f:
            yaml.safe_dump(
                data,
                f,
                sort_keys=False,
                allow_unicode=True,
                width=80,
            )

        print(f"Wrote {len(questions)} questions to {output_yaml_path}")
        if image_count:
            print(f"Copied {image_count} image(s) to {images_dir}")


def convert_subcommand(args):
    """CLI subcommand wrapper for :func:`convert`."""
    convert(args.input_zip, args.output_yaml)
    return 0


def main():
    if len(sys.argv) != 3:
        print("Usage: python zybooks_zip_to_yaml.py export.zip output.yaml")
        sys.exit(1)

    convert(sys.argv[1], sys.argv[2])


if __name__ == "__main__":
    main()
